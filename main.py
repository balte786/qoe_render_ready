
from fastapi import FastAPI, UploadFile, File, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse
import pandas as pd
import matplotlib.pyplot as plt
import io
import openai
import os
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openai import OpenAI

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

openai.api_key = os.getenv("OPENAI_API_KEY")

qoe_cache = {}

def build_prompt(prompt_type: str, data: str) -> str:
    templates = {
        "executive_summary": f"Create an executive summary for a QoE report using this data: {data}",
        "revenue_trends": f"Analyze revenue trends in this monthly income statement: {data}",
        "addbacks": f"Identify one-time or non-recurring expenses for EBITDA adjustments: {data}",
        "working_capital": f"Calculate and normalize working capital from this data, highlighting changes in AR, Inventory, AP: {data}",
    }
    return templates.get(prompt_type, "Invalid prompt type")

@app.post("/upload")
async def upload_file(file: UploadFile = File(...), type: str = Query("pnl")):
    contents = await file.read()
    df = pd.read_excel(io.BytesIO(contents))
    if type == "balance_sheet":
        qoe_cache["balance_sheet"] = df
        return {"message": "Balance Sheet uploaded."}
    else:
        qoe_cache["pnl_df"] = df
        summary = df.describe().to_dict()
        qoe_cache['data'] = df.to_dict(orient='records')
        return {"message": "P&L file received", "summary": summary, "data": qoe_cache['data']}

client = OpenAI()

@app.post("/generate_qoe")
async def generate_qoe(payload: dict):
    prompt_type = payload.get("type", "executive_summary")
    data = payload.get("financial_summary", "")
    if prompt_type == "working_capital" and "balance_sheet" in qoe_cache:
        bs_preview = qoe_cache["balance_sheet"].to_string(index=False)
        data += f"\n\nBalance Sheet Data:\n{bs_preview}"
    
    prompt = build_prompt(prompt_type, data)

    response = client.chat.completions.create(
        model="gpt-4-turbo",
        messages=[
            {"role": "system", "content": "You are a financial due diligence analyst."},
            {"role": "user", "content": prompt}
        ]
    )

    content = response.choices[0].message.content  # <- Corrected here

    if 'qoe_report' not in qoe_cache:
        qoe_cache['qoe_report'] = {}
    qoe_cache['qoe_report'][prompt_type] = content

    return {"qoe_summary": content}

@app.get("/export_docx")
async def export_docx():
    doc = Document()
    doc.add_heading("Acquisight QoE Report", 0)

    if 'pnl_df' in qoe_cache:
        doc.add_heading("P&L Table Preview", level=1)
        df = qoe_cache['pnl_df'].head()
        table = doc.add_table(rows=1, cols=len(df.columns))
        table.style = 'Light Grid Accent 1'
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr_cells[i].text = str(col)
        for row in df.values.tolist():
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)

    if 'balance_sheet' in qoe_cache:
        doc.add_heading("Balance Sheet Preview", level=1)
        bs_df = qoe_cache["balance_sheet"].head()
        table = doc.add_table(rows=1, cols=len(bs_df.columns))
        table.style = 'Light Grid Accent 2'
        hdr_cells = table.rows[0].cells
        for i, col in enumerate(bs_df.columns):
            hdr_cells[i].text = str(col)
        for row in bs_df.values.tolist():
            row_cells = table.add_row().cells
            for i, cell in enumerate(row):
                row_cells[i].text = str(cell)

    if 'qoe_report' in qoe_cache:
        for section, text in qoe_cache['qoe_report'].items():
            doc.add_heading(section.replace("_", " ").title(), level=1)
            para = doc.add_paragraph(text)
            para.style.font.size = Pt(11)

    path = "qoe_report.docx"
    doc.save(path)
    return FileResponse(path, media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document', filename=path)

@app.get("/revenue_chart")
async def revenue_chart():
    if 'pnl_df' not in qoe_cache:
        return {"error": "No data uploaded."}
    df = qoe_cache['pnl_df']
    if 'Date' not in df.columns or 'Revenue' not in df.columns:
        return {"error": "Missing 'Date' or 'Revenue' columns."}
    df['Date'] = pd.to_datetime(df['Date'])
    df = df.sort_values('Date')
    plt.figure(figsize=(10, 5))
    plt.plot(df['Date'], df['Revenue'], marker='o')
    plt.title('Revenue Trend Over Time')
    plt.xlabel('Date')
    plt.ylabel('Revenue')
    plt.grid(True)
    path = "revenue_chart.png"
    plt.savefig(path)
    plt.close()
    return FileResponse(path, media_type="image/png", filename=path)



if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
